#!/usr/bin/env python3
"""
read_output.py — 读取填充后的 XLSX/DOCX 文件，输出所有非空单元格内容
用于 AI 验证阶段（Step 4）读取输出文件。

用法:
    python3 read_output.py <file_path>
"""

import sys


def read_xlsx(file_path: str):
    """读取 XLSX 文件，打印所有非空单元格"""
    import openpyxl
    wb = openpyxl.load_workbook(file_path, data_only=True)
    ws = wb.active
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=10, values_only=False):
        cells = []
        for cell in row:
            if cell.value is not None:
                cells.append(f"{cell.coordinate}={cell.value}")
        if cells:
            print(f"Row {row[0].row}: {'  |  '.join(cells)}")
    wb.close()


def read_docx(file_path: str):
    """读取 DOCX 文件，打印段落和表格内容"""
    from docx import Document
    doc = Document(file_path)

    print("=== 段落 ===")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"Para {i}: {text}")

    print("\n=== 表格 ===")
    for ti, table in enumerate(doc.tables):
        print(f"\nTable {ti} ({len(table.rows)} rows x {len(table.columns)} cols):")
        for ri, row in enumerate(table.rows):
            cells = []
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    cells.append(f"C{ci}={text}")
            if cells:
                print(f"  Row {ri}: {'  |  '.join(cells)}")


def main():
    if len(sys.argv) < 2:
        print("用法: python3 read_output.py <file_path>", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]
    ext = file_path.lower().rsplit('.', 1)[-1] if '.' in file_path else ''

    if ext == 'xlsx':
        read_xlsx(file_path)
    elif ext == 'docx':
        read_docx(file_path)
    else:
        print(f"不支持的文件格式: .{ext}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
