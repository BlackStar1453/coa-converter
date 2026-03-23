#!/usr/bin/env python3
"""
DOCX Template Filler
基于 TemplateLayout 动态填充 DOCX 模板。
"""

import os
import re
import sys
import shutil
import logging
from typing import Dict, Optional

from docx import Document
from docx.shared import RGBColor

from template_detector import TemplateLayout
from coa_converter import COAData

logger = logging.getLogger(__name__)


def fill_docx(coa: COAData, layout: TemplateLayout, template_path: str, output_path: str):
    """根据 layout 动态填充 DOCX 模板"""
    logger.info(f'[DOCX填充] 模板类型: {layout.template_type}')
    logger.info(f'[DOCX填充] 模板: {template_path} → {output_path}')

    shutil.copy2(template_path, output_path)
    # 确保输出文件可写（模板可能是只读的）
    if sys.platform != 'win32':
        os.chmod(output_path, 0o644)
    doc = Document(output_path)

    product_name = coa.header.get('product_name', '')
    logger.info(f'[DOCX填充] 产品名: {product_name}')

    if layout.template_type == 'cs':
        _fill_cs(doc, coa, layout)
    elif layout.template_type == 'nutrition':
        _fill_nutrition(doc, coa, layout)
    elif layout.template_type == 'sds':
        _fill_sds(doc, coa, layout)
    elif layout.template_type in ('composition_powder', 'composition_standard'):
        _fill_composition(doc, coa, layout)
    else:
        logger.warning(f'[DOCX填充] 未知模板类型: {layout.template_type}，仅替换产品名')
        _replace_product_name_in_paragraphs(doc, product_name)

    doc.save(output_path)
    logger.info(f'[DOCX填充] 文件已保存: {output_path}')


# ============ 段落文本替换工具 ============

def _replace_in_run(run, old_text: str, new_text: str) -> bool:
    """在单个 run 中替换文本，保持格式"""
    if old_text in run.text:
        run.text = run.text.replace(old_text, new_text)
        return True
    return False


def _replace_in_paragraph(paragraph, old_text: str, new_text: str) -> bool:
    """在段落中替换文本（先尝试 run 级别，再尝试跨 run 拼接）"""
    # 方式1: 单个 run 内替换
    for run in paragraph.runs:
        if _replace_in_run(run, old_text, new_text):
            return True

    # 方式2: 跨 run 拼接替换（文本可能被拆分到多个 run 中）
    full_text = paragraph.text
    if old_text in full_text:
        # 找到在哪些 run 中
        pos = full_text.find(old_text)

        current_pos = 0
        for idx, run in enumerate(paragraph.runs):
            run_end = current_pos + len(run.text)
            if current_pos <= pos < run_end:
                # 在这个 run 中开始替换
                offset = pos - current_pos
                remaining = run.text[offset:]
                if old_text.startswith(remaining):
                    # 替换可能跨多个 run
                    run.text = run.text[:offset] + new_text
                    # 清除后续 run 中属于 old_text 的部分
                    chars_consumed = len(remaining)
                    for next_run in paragraph.runs[idx + 1:]:
                        if chars_consumed >= len(old_text):
                            break
                        chars_to_remove = min(len(next_run.text), len(old_text) - chars_consumed)
                        next_run.text = next_run.text[chars_to_remove:]
                        chars_consumed += chars_to_remove
                    return True
            current_pos = run_end

    return False


def _replace_product_name_in_paragraphs(doc, product_name: str):
    """在所有段落中查找占位产品名并替换"""
    if not product_name:
        return

    for para in doc.paragraphs:
        text = para.text
        # 查找 "our product XXX is:" 或 "our product, XXX," 模式
        # 替换占位产品名（通常是模板中的示例名称）
        m = re.search(r'our product[,\s]+(.+?)(?:\s+is[:\s]|\s*,)', text)
        if m:
            old_name = m.group(1).strip()
            if old_name and old_name != product_name:
                _replace_in_paragraph(para, old_name, product_name)
                logger.info(f'[DOCX填充] 替换产品名: "{old_name}" → "{product_name}"')


def _replace_product_name_in_tables(doc, product_name: str):
    """在所有表格中查找 product name 单元格并替换"""
    if not product_name:
        return

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell_lower = cell.text.strip().lower()
                if 'product name' in cell_lower:
                    # 找到标签单元格，值通常在同行下一列
                    if ci + 1 < len(row.cells):
                        val_cell = row.cells[ci + 1]
                        old_val = val_cell.text.strip()
                        if old_val and old_val != product_name:
                            # 替换值单元格中的文本
                            for run in val_cell.paragraphs[0].runs:
                                if run.text.strip():
                                    run.text = product_name
                                    logger.info(f'[DOCX填充] Table[{ti}] 产品名: "{old_val}" → "{product_name}"')
                                    break
                    break


# ============ CS (Combined Statement) 模板 ============

def _fill_cs(doc, coa: COAData, layout: TemplateLayout):
    """填充 Combined Statement 模板 — 替换产品名 + 原产国"""
    product_name = coa.header.get('product_name', '')
    country = coa.header.get('country', '')

    _replace_product_name_in_paragraphs(doc, product_name)

    # 替换 Country of Origin
    if country:
        for para in doc.paragraphs:
            text = para.text
            m = re.match(r'(Country of Origin\s*[-–—]\s*)(.*)', text)
            if m:
                old_country = m.group(2).strip()
                if old_country and old_country != country:
                    _replace_in_paragraph(para, old_country, country)
                    logger.info(f'[DOCX填充-CS] Country of Origin: "{old_country}" → "{country}"')
                break

    logger.info('[DOCX填充-CS] Combined Statement 已填充产品名和原产国')


# ============ Nutrition 模板 ============

def _fill_nutrition(doc, coa: COAData, layout: TemplateLayout):
    """填充 Nutrition Info 模板 — 替换产品名 + 填充营养数据表"""
    product_name = coa.header.get('product_name', '')
    _replace_product_name_in_paragraphs(doc, product_name)

    # 营养数据通常不来自 COA PDF，COA 主要包含检测数据
    # 如果 coa 中有营养数据，可以填充到表格中
    logger.info('[DOCX填充-Nutrition] 产品名已替换，营养数据表保持模板默认值')


# ============ SDS (Safety Data Sheet) 模板 ============

def _fill_sds(doc, coa: COAData, layout: TemplateLayout):
    """填充 Safety Data Sheet 模板 — 替换产品名和物料名"""
    product_name = coa.header.get('product_name', '')
    botanical_name = coa.header.get('botanical_name', '')

    _replace_product_name_in_tables(doc, product_name)

    # 替换 Material Name（如果有植物学名）
    if botanical_name:
        for table in doc.tables:
            for row in table.rows:
                for ci, cell in enumerate(row.cells):
                    if 'material name' in cell.text.strip().lower():
                        if ci + 1 < len(row.cells):
                            val_cell = row.cells[ci + 1]
                            old_val = val_cell.text.strip()
                            if old_val != botanical_name:
                                for run in val_cell.paragraphs[0].runs:
                                    if run.text.strip():
                                        run.text = botanical_name
                                        logger.info(f'[DOCX填充-SDS] Material Name: "{old_val}" → "{botanical_name}"')
                                        break
                        break

    # 去除表格中的黄色高亮标记（模板使用高亮标注需填充的单元格）
    highlight_cleared = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.highlight_color is not None:
                            run.font.highlight_color = None
                            highlight_cleared += 1
    if highlight_cleared:
        logger.info(f'[DOCX填充-SDS] 已去除 {highlight_cleared} 个高亮标记')

    logger.info('[DOCX填充-SDS] SDS 模板已填充产品名和物料名，其余数据保持模板默认值')


# ============ Composition Statement 模板 ============

def _fill_composition(doc, coa: COAData, layout: TemplateLayout):
    """填充 Composition Statement 模板 — 替换段落和表格中的产品名"""
    product_name = coa.header.get('product_name', '')
    if not product_name:
        logger.warning('[DOCX填充-Composition] 产品名为空，跳过填充')
        return

    # 1. 替换段落中 "our product, XXX, is" 的产品名
    #    模板中占位符可能是空格或旧示例名，通用 regex 匹配不到，需要专门处理
    #    使用包含上下文的唯一匹配串来避免替换到段落中其他位置的同名文本
    for para in doc.paragraphs:
        full_text = para.text
        m = re.search(r'our product,(.+?),\s*is', full_text)
        if m:
            old_segment = m.group(1)  # 两个逗号之间的内容（可能是空格或旧名称）
            if old_segment.strip() == product_name:
                continue
            # 使用包含上下文的完整匹配串进行替换，避免歧义
            old_full = m.group(0)  # e.g. "our product,         ,is"
            new_full = f'our product, {product_name} ,is'
            if _replace_in_paragraph(para, old_full, new_full):
                logger.info(f'[DOCX填充-Composition] 段落产品名: "{old_segment.strip()}" → "{product_name}"')
            else:
                # 回退：直接重写 runs，但也使用完整上下文
                _rebuild_paragraph_with_product(para, old_full, new_full)
                logger.info(f'[DOCX填充-Composition] 段落产品名(重建): "{old_segment.strip()}" → "{product_name}"')

    # 2. 填充表格
    for ti, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue
        row1 = table.rows[1]

        # Row[1] Col[0] = 产品名
        if len(row1.cells) >= 1:
            cell0 = row1.cells[0]
            old_val = cell0.text.strip()
            if old_val != product_name:
                _set_cell_text(cell0, product_name)
                logger.info(f'[DOCX填充-Composition] Table[{ti}] R1C0: "{old_val}" → "{product_name}"')

        # Row[1] Col[1] = Content value（在百分比后嵌入产品名）
        if len(row1.cells) >= 2:
            cell1 = row1.cells[1]
            _fill_composition_content_cell(cell1, product_name, ti)

    # 3. 清除红色字体标记（模板使用 #EE0000 红字作为格式提示，不应出现在输出中）
    red_color = RGBColor(0xEE, 0x00, 0x00)
    red_cleared = 0
    for para in doc.paragraphs:
        # 跳过标题段落，保留比例文本（如 "10:1", "No Maltodextrin Used"）
        if para.style.name == 'Title':
            continue
        for run in para.runs:
            if run.font.color and run.font.color.rgb == red_color:
                run.text = ''
                red_cleared += 1
    if red_cleared:
        logger.info(f'[DOCX填充-Composition] 已清除 {red_cleared} 个红色标记 run')

    logger.info('[DOCX填充-Composition] Composition Statement 已填充产品名')


def _set_cell_text(cell, text: str):
    """设置单元格文本，保留第一个 run 的格式，清除多余段落"""
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
        for run in cell.paragraphs[0].runs[1:]:
            run.text = ''
    elif cell.paragraphs:
        cell.paragraphs[0].text = text
    # 清除单元格中多余的段落文本（保留段落结构但清空内容）
    for para in cell.paragraphs[1:]:
        for run in para.runs:
            run.text = ''


def _rebuild_paragraph_with_product(para, old_segment: str, new_segment: str):
    """当 _replace_in_paragraph 失败时，通过重建 runs 替换段落中的文本片段"""
    full_text = para.text
    new_text = full_text.replace(old_segment, new_segment)
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''


def _fill_composition_content_cell(cell, product_name: str, table_index: int):
    """填充 Composition 表格的 Content value 单元格。

    将 "88% " → "88% PRODUCT"、"100%" → "100% PRODUCT"，
    保持 "12% Maltodextrin" 不变。
    """
    for para in cell.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 跳过 Maltodextrin 行
        if 'maltodextrin' in text.lower():
            continue
        # 匹配 "NN% optional_text" 格式
        m = re.match(r'(\d+%)\s*(.*)', text)
        if m:
            percentage = m.group(1)
            old_rest = m.group(2).strip()
            new_text = f'{percentage} {product_name}'
            if old_rest == product_name:
                break  # 已正确
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ''
            logger.info(f'[DOCX填充-Composition] Table[{table_index}] Content: "{text}" → "{new_text}"')
            break  # 只处理第一行（产品行）


# ============ DOCX 验证 ============

def verify_docx_output(coa: COAData, layout: TemplateLayout, output_path: str) -> dict:
    """验证 DOCX 输出文件的填充正确性"""
    doc = Document(output_path)

    total = 0
    passed = 0
    failed = 0
    details = []

    product_name = coa.header.get('product_name', '')
    country = coa.header.get('country', '')
    botanical_name = coa.header.get('botanical_name', '')

    all_para_text = '\n'.join(p.text for p in doc.paragraphs)
    all_table_text = '\n'.join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    full_text = all_para_text + '\n' + all_table_text

    # 通用：检查产品名是否出现在文档中
    if product_name:
        total += 1
        if product_name in full_text:
            passed += 1
        else:
            failed += 1
            details.append(f'产品名 "{product_name}" 未出现在文档中')

    # CS 特有：检查 Country of Origin
    if layout.template_type == 'cs' and country:
        total += 1
        if country in all_para_text:
            passed += 1
        else:
            failed += 1
            details.append(f'Country of Origin "{country}" 未出现在文档中')

    # SDS 特有：检查 Material Name
    if layout.template_type == 'sds' and botanical_name:
        total += 1
        if botanical_name in all_table_text:
            passed += 1
        else:
            failed += 1
            details.append(f'Material Name "{botanical_name}" 未出现在表格中')

    # Composition 特有：检查每个表格 Row[1] 中是否包含产品名
    if layout.template_type in ('composition_powder', 'composition_standard') and product_name:
        for ti, table in enumerate(doc.tables):
            if len(table.rows) >= 2:
                total += 1
                row1_text = ' '.join(cell.text for cell in table.rows[1].cells)
                if product_name in row1_text:
                    passed += 1
                else:
                    failed += 1
                    details.append(f'Table[{ti}] Row[1] 未包含产品名')

    # 检查无残留模板默认值
    example_names = ['Ashwagandha', 'Pumpkin Seed Protein Powder', 'Organic Moringa Extract']
    for ex_name in example_names:
        if ex_name in full_text and product_name and ex_name not in product_name:
            total += 1
            failed += 1
            details.append(f'残留模板示例值: "{ex_name}"')

    accuracy = passed / total if total > 0 else 1.0
    result = {
        'total': total,
        'passed': passed,
        'failed': failed,
        'accuracy': accuracy,
        'details': details,
    }

    if failed > 0:
        logger.warning(f'[DOCX验证] {failed}/{total} 项不通过:')
        for d in details:
            logger.warning(f'  - {d}')
    else:
        logger.info(f'[DOCX验证] 全部通过 ({passed}/{total})，正确率: {accuracy:.1%}')

    return result
