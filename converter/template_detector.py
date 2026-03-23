#!/usr/bin/env python3
"""
Template Layout Detector
自动检测 XLSX/DOCX 模板的字段布局，返回 TemplateLayout 数据结构。
"""

import re
import logging
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple

import openpyxl

logger = logging.getLogger(__name__)

# 复用 coa_converter 中的别名映射
from coa_converter import HEADER_FIELD_ALIASES, ITEM_NAME_NORMALIZE, normalize_item_name


# ============ 数据结构 ============

@dataclass
class Position:
    """通用位置描述"""
    # XLSX: row, col (1-based)
    # DOCX: para_index or (table_index, row_index, col_index)
    row: Optional[int] = None
    col: Optional[int] = None
    para_index: Optional[int] = None
    table_index: Optional[int] = None
    cell_ref: Optional[str] = None  # e.g. "C4" for XLSX


@dataclass
class FieldMapping:
    """单个字段的位置映射"""
    label_pos: Optional[Position] = None
    value_pos: Optional[Position] = None


@dataclass
class TemplateLayout:
    """模板布局检测结果"""
    format: str  # "xlsx" or "docx"
    template_type: str  # "coa_assay", "coa_powder", "coa_ratio", "allergen", "flowchart",
                        # "composition_powder", "composition_standard", "cs", "nutrition", "sds"

    # XLSX 字段
    header_fields: Dict[str, FieldMapping] = field(default_factory=dict)
    table_rows: Dict[str, int] = field(default_factory=dict)       # item_key → row number
    data_columns: Dict[str, int] = field(default_factory=dict)     # "spec"/"result"/"method" → col number
    label_col: int = 1             # 标签列号 (1=A)
    packing_position: Optional[Position] = None
    storage_position: Optional[Position] = None
    data_header_row: Optional[int] = None  # "Items of Analysis" 行号

    # DOCX 字段
    product_name_positions: List[Position] = field(default_factory=list)
    docx_tables: List[Dict] = field(default_factory=list)  # 每个表格的字段映射


# ============ XLSX 检测 ============

# 用于匹配数据表头行的关键词
_DATA_HEADER_KW = {'items of analysis', 'items', 'specification', 'result', 'test method', 'parameter'}

# 用于识别 section header 的模式
_SECTION_HEADERS = {
    'analytical data': 'analytical',
    'analytical': 'analytical',
    'microbiological test': 'microbiology',
    'microbiological': 'microbiology',
    'additional information': 'additional',
    'packing and storage': 'packing',
    'packing & storage': 'packing',
}

# Allergen 模板特有的标签列标识
_ALLERGEN_MARKERS = {'cereals', 'crustaceans', 'eggs', 'fish', 'peanuts', 'soybeans', 'milk'}


def detect_xlsx_layout(template_path: str) -> TemplateLayout:
    """检测 XLSX 模板的字段布局"""
    logger.info(f'[检测-XLSX] 扫描模板: {template_path}')

    wb = openpyxl.load_workbook(template_path, data_only=True)
    ws = wb.active
    layout = TemplateLayout(format="xlsx", template_type="unknown")

    # Phase 1: 扫描所有单元格，构建文本索引
    cell_texts = {}  # (row, col) → text
    for row in ws.iter_rows(min_row=1, max_row=50, max_col=10, values_only=False):
        for cell in row:
            if cell.value is not None:
                text = str(cell.value).strip()
                if text:
                    cell_texts[(cell.row, cell.column)] = text

    # Phase 2: 判断模板类型
    layout.template_type = _classify_xlsx_template(cell_texts)
    logger.info(f'[检测-XLSX] 模板类型: {layout.template_type}')

    if layout.template_type == 'allergen':
        _detect_allergen_layout(cell_texts, layout)
    elif layout.template_type == 'flowchart':
        _detect_flowchart_layout(cell_texts, layout)
    else:
        _detect_coa_layout(cell_texts, layout)

    wb.close()

    logger.info(f'[检测-XLSX] 完成 - header_fields: {len(layout.header_fields)}, '
                f'table_rows: {len(layout.table_rows)}, data_columns: {layout.data_columns}')
    return layout


def _classify_xlsx_template(cell_texts: Dict[Tuple[int, int], str]) -> str:
    """根据内容判断 XLSX 模板类型"""
    all_text_lower = ' '.join(v.lower() for v in cell_texts.values())

    if 'allergen statement' in all_text_lower:
        return 'allergen'
    if 'flow chart' in all_text_lower:
        return 'flowchart'

    # COA 模板：区分 Assay / Powder / Ratio
    # 检查是否有 Assay 行、Ratio 行、还是直接 Analytical Data
    for (r, c), text in cell_texts.items():
        text_lower = text.lower().strip()
        if text_lower == 'assay':
            return 'coa_assay'
        if text_lower == 'ratio':
            return 'coa_ratio'

    # 没有 Assay/Ratio 行 → Powder 模板
    if 'certificate of analysis' in all_text_lower:
        return 'coa_powder'

    return 'coa_assay'  # 默认


def _detect_coa_layout(cell_texts: Dict[Tuple[int, int], str], layout: TemplateLayout):
    """检测 COA 类型 XLSX 模板（Assay/Powder/Ratio）的布局"""
    layout.label_col = 1  # Column A

    # 1. 扫描头部字段（Row 1-8）
    for (r, c), text in cell_texts.items():
        if r > 8:
            continue
        text_lower = text.lower().strip()
        # 检查是否是已知的头部标签
        for alias, std_key in HEADER_FIELD_ALIASES.items():
            if alias in text_lower or text_lower in alias:
                # 标签在 (r, c)，值在同一行的右侧
                value_col = c + 2 if c == 1 else c + 1  # A→C 或 E→F
                # 查找实际有值的列
                for vc in range(c + 1, c + 4):
                    if (r, vc) in cell_texts:
                        value_col = vc
                        break
                layout.header_fields[std_key] = FieldMapping(
                    label_pos=Position(row=r, col=c),
                    value_pos=Position(row=r, col=value_col,
                                       cell_ref=f'{_col_letter(value_col)}{r}')
                )
                break

    # 2. 找到数据表头行（"Items of Analysis"）
    for (r, c), text in cell_texts.items():
        text_lower = text.lower().strip()
        if 'items of analysis' in text_lower or text_lower == 'items' or text_lower == 'parameter':
            layout.data_header_row = r
            break

    if not layout.data_header_row:
        logger.warning('[检测-XLSX] 未找到数据表头行')
        return

    # 3. 检测数据列号（Specification / Result / Test Method）
    dhr = layout.data_header_row
    for (r, c), text in cell_texts.items():
        if r != dhr:
            continue
        text_lower = text.lower().strip()
        if 'specification' in text_lower:
            layout.data_columns['spec'] = c
        elif 'result' in text_lower:
            layout.data_columns['result'] = c
        elif 'test method' in text_lower or 'method' in text_lower:
            layout.data_columns['method'] = c

    # 默认列号（如果未检测到）
    layout.data_columns.setdefault('spec', 3)    # C
    layout.data_columns.setdefault('result', 5)   # E
    layout.data_columns.setdefault('method', 6)   # F

    # 4. 扫描数据行（表头之后），匹配检测项标签 → 行号
    for (r, c), text in cell_texts.items():
        if r <= layout.data_header_row or c != layout.label_col:
            continue
        text_lower = text.lower().strip()

        # Packing and Storage（必须在 section header 跳过之前检测）
        if 'packing' in text_lower or ('storage' in text_lower and 'packing' not in text_lower):
            spec_col = layout.data_columns.get('spec', 3)
            # label 在 A 列，packing 数据在同行 C 列（spec_col）
            packing_row = r
            layout.packing_position = Position(row=packing_row, col=spec_col,
                                                cell_ref=f'{_col_letter(spec_col)}{packing_row}')
            # Storage 在下一行
            storage_row = packing_row + 1
            if any((storage_row, cc) in cell_texts for cc in range(1, 10)):
                layout.storage_position = Position(row=storage_row, col=spec_col,
                                                    cell_ref=f'{_col_letter(spec_col)}{storage_row}')
            logger.info(f'[检测-XLSX-COA] packing_position: {layout.packing_position.cell_ref}, '
                        f'storage_position: {layout.storage_position.cell_ref if layout.storage_position else "None"}')
            continue

        # 检查 section headers
        if text_lower in _SECTION_HEADERS:
            continue

        # 尝试匹配检测项
        item_key = normalize_item_name(text)
        if item_key:
            layout.table_rows[item_key] = r
            continue

    logger.info(f'[检测-XLSX-COA] table_rows: {layout.table_rows}')


def _detect_allergen_layout(cell_texts: Dict[Tuple[int, int], str], layout: TemplateLayout):
    """检测 Allergen 模板布局"""
    layout.label_col = 2  # Column B

    # 找 Product Name 行
    for (r, c), text in cell_texts.items():
        if 'product name' in text.lower():
            layout.header_fields['product_name'] = FieldMapping(
                label_pos=Position(row=r, col=c),
                value_pos=Position(row=r, col=c, cell_ref=f'{_col_letter(c)}{r}')
            )
            break

    # 找表头行（Items / Contain / Absent ...）
    for (r, c), text in cell_texts.items():
        text_lower = text.lower().strip()
        if text_lower == 'items':
            layout.data_header_row = r
            # 检测数据列
            for (r2, c2), t2 in cell_texts.items():
                if r2 != r:
                    continue
                t2l = t2.lower().strip()
                if 'contain' in t2l:
                    layout.data_columns['contain'] = c2
                elif 'absent' in t2l:
                    layout.data_columns['absent'] = c2
                elif 'same production line' in t2l:
                    layout.data_columns['production_line'] = c2
                elif 'same production facility' in t2l:
                    layout.data_columns['production_facility'] = c2
                elif 'comment' in t2l:
                    layout.data_columns['comments'] = c2
            break

    # 扫描过敏原项目行
    if layout.data_header_row:
        for (r, c), text in cell_texts.items():
            if r <= layout.data_header_row or c != layout.label_col:
                continue
            if text.strip():
                # 用标签文本的前几个词作为key
                key = re.sub(r'[^a-z\s]', '', text.lower().strip()).split()[0] if text.strip() else ''
                if key:
                    layout.table_rows[key] = r


def _detect_flowchart_layout(cell_texts: Dict[Tuple[int, int], str], layout: TemplateLayout):
    """检测 Flow Chart 模板布局"""
    layout.label_col = 1  # Column A
    layout.data_columns['value'] = 3  # Column C

    # 找标题行（含 "Flow Chart"）
    for (r, c), text in cell_texts.items():
        if 'flow chart' in text.lower():
            layout.header_fields['title'] = FieldMapping(
                label_pos=Position(row=r, col=c),
                value_pos=Position(row=r, col=c, cell_ref=f'{_col_letter(c)}{r}')
            )
            break


# ============ DOCX 检测 ============

def detect_docx_layout(template_path: str) -> TemplateLayout:
    """检测 DOCX 模板的字段布局"""
    from docx import Document

    logger.info(f'[检测-DOCX] 扫描模板: {template_path}')

    doc = Document(template_path)
    layout = TemplateLayout(format="docx", template_type="unknown")

    # 判断类型
    all_para_text = '\n'.join(p.text for p in doc.paragraphs).lower()
    num_tables = len(doc.tables)

    if 'combined statement' in all_para_text:
        layout.template_type = 'cs'
    elif 'safety data sheet' in all_para_text:
        layout.template_type = 'sds'
    elif 'standardized material' in all_para_text:
        layout.template_type = 'composition_standard'
    elif 'composition statement' in all_para_text:
        layout.template_type = 'composition_powder'
    elif 'nutrition statement' in all_para_text or 'nutrition info' in all_para_text:
        layout.template_type = 'nutrition'
    else:
        layout.template_type = 'docx_generic'

    logger.info(f'[检测-DOCX] 模板类型: {layout.template_type}, '
                f'段落数: {len(doc.paragraphs)}, 表格数: {num_tables}')

    # 扫描段落，定位产品名和 Key-In Nutrition
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 查找 "Key-In Nutrition" 出现位置
        if 'key-in nutrition' in text.lower():
            layout.product_name_positions.append(
                Position(para_index=i)
            )

    # 扫描表格结构
    for ti, table in enumerate(doc.tables):
        table_info = {
            'table_index': ti,
            'rows': len(table.rows),
            'cols': len(table.columns),
            'fields': {}
        }

        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                cell_lower = cell_text.lower()

                # 检查是否是已知标签
                if 'product name' in cell_lower:
                    table_info['fields']['product_name'] = (ri, ci)
                elif 'content value' in cell_lower:
                    table_info['fields']['content_value'] = (ri, ci)
                elif cell_lower == 'item':
                    table_info['fields']['item_header'] = (ri, ci)
                elif 'result' in cell_lower:
                    table_info['fields']['result_header'] = (ri, ci)

                # SDS 特有标签
                for alias, std_key in HEADER_FIELD_ALIASES.items():
                    if alias in cell_lower:
                        table_info['fields'][std_key] = (ri, ci)
                        break

        layout.docx_tables.append(table_info)

    logger.info(f'[检测-DOCX] 完成 - product_name_positions: {len(layout.product_name_positions)}, '
                f'tables: {len(layout.docx_tables)}')
    return layout


# ============ 统一入口 ============

def detect_template_layout(template_path: str) -> TemplateLayout:
    """自动检测模板格式和布局"""
    ext = template_path.lower().rsplit('.', 1)[-1] if '.' in template_path else ''

    if ext == 'xlsx':
        return detect_xlsx_layout(template_path)
    elif ext == 'docx':
        return detect_docx_layout(template_path)
    else:
        raise ValueError(f'不支持的模板格式: .{ext} (仅支持 .xlsx 和 .docx)')


# ============ 工具函数 ============

def _col_letter(col_num: int) -> str:
    """列号转字母 (1→A, 2→B, ... 26→Z)"""
    result = ''
    while col_num > 0:
        col_num, remainder = divmod(col_num - 1, 26)
        result = chr(65 + remainder) + result
    return result
