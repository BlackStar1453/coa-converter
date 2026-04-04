#!/usr/bin/env python3
"""
fill_template.py — JSON 驱动的模板填充工具
AI-First 架构的核心：接受 AI 生成的 JSON 数据，机械填充 XLSX/DOCX 模板。

用法:
    python3 fill_template.py <json_data_path> <template_path> <output_path>

JSON 格式参见 coa_schema.json
"""

import os
import re
import sys
import json
import shutil
import logging
import zipfile
import tempfile
from collections import defaultdict
from typing import Dict, List, Optional

import openpyxl

from template_detector import (
    detect_template_layout, TemplateLayout, _col_letter
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s'
)
logger = logging.getLogger(__name__)


# ============ 主入口 ============

def fill_from_json(data: dict, template_path: str, output_path: str):
    """根据 JSON 数据填充模板"""
    layout = detect_template_layout(template_path)
    logger.info(f'[填充] 模板类型: {layout.template_type}, 格式: {layout.format}')

    if layout.format == 'xlsx':
        _fill_xlsx(data, layout, template_path, output_path)
    elif layout.format == 'docx':
        _fill_docx(data, layout, template_path, output_path)
    else:
        raise ValueError(f'不支持的模板格式: {layout.format}')


# ============ XLSX 填充 ============

def _fill_xlsx(data: dict, layout: TemplateLayout, template_path: str, output_path: str):
    """填充 XLSX 模板"""
    shutil.copy2(template_path, output_path)
    if sys.platform != 'win32':
        os.chmod(output_path, 0o644)

    # Flow Chart 使用 XML 直接操作
    if layout.template_type == 'flowchart':
        _fill_flowchart_xml(data, layout, output_path)
        return

    wb = openpyxl.load_workbook(output_path)
    ws = wb.active

    if layout.template_type in ('coa_assay', 'coa_powder', 'coa_ratio'):
        _fill_coa_xlsx(ws, data, layout)
    elif layout.template_type == 'allergen':
        _fill_allergen_xlsx(ws, data, layout)
    else:
        logger.warning(f'[填充] 未知模板类型 {layout.template_type}，按 COA 方式填充')
        _fill_coa_xlsx(ws, data, layout)

    wb.save(output_path)
    logger.info(f'[填充] XLSX 已保存: {output_path}')


def _fill_coa_xlsx(ws, data: dict, layout: TemplateLayout):
    """填充 COA 类型 XLSX 模板"""
    header = data.get("header", {})
    assay = data.get("assay")
    test_items = data.get("test_items", [])
    mapping = data.get("template_mapping", [])
    packing = data.get("packing", "")
    storage = data.get("storage", "")

    spec_col = layout.data_columns.get('spec', 3)
    result_col = layout.data_columns.get('result', 5)
    method_col = layout.data_columns.get('method', 6)
    label_col = getattr(layout, 'label_col', 1)

    # --- 填充头部字段 ---
    for std_key, field_mapping in layout.header_fields.items():
        value = header.get(std_key, "")
        ref = field_mapping.value_pos.cell_ref
        if ref:
            ws[ref] = value
        else:
            ws.cell(row=field_mapping.value_pos.row,
                    column=field_mapping.value_pos.col, value=value)
        if value:
            logger.info(f'[填充-头部] {std_key} = {value}')
        elif std_key in ('product_name', 'batch_number'):
            logger.warning(f'[填充-头部] 缺少字段: {std_key}')

    # --- 填充 Assay/Ratio 行 ---
    assay_row = layout.table_rows.get('assay')
    if assay_row:
        if assay:
            ws.cell(row=assay_row, column=spec_col, value=assay.get("specification", ""))
            ws.cell(row=assay_row, column=result_col, value=assay.get("result", ""))
            ws.cell(row=assay_row, column=method_col, value=assay.get("method", ""))
            logger.info(f'[填充-Assay] Row {assay_row}: {assay}')
        else:
            ws.cell(row=assay_row, column=spec_col, value="")
            ws.cell(row=assay_row, column=result_col, value="")
            ws.cell(row=assay_row, column=method_col, value="")

    # --- 构建映射索引 ---
    # 已映射到模板行的 item_key 集合
    mapped_keys = set()
    # 额外项列表（按 mapping 顺序）
    extra_items = []
    # template_row_label → test_item
    label_to_item = {}

    for m in mapping:
        idx = m.get("test_item_index")
        if idx is None or idx >= len(test_items):
            continue
        item = test_items[idx]
        if m.get("is_extra"):
            extra_items.append({
                "item": item,
                "insert_after_label": m.get("insert_after_label")
            })
        else:
            label = m.get("template_row_label", "")
            if label:
                label_to_item[label.lower().strip()] = item

    # --- 填充模板中的数据行 ---
    # 遍历模板行，从 label_to_item 中查找对应数据
    for item_key, row_num in layout.table_rows.items():
        if item_key == 'assay':
            continue
        # 尝试从 template_mapping 中查找（通过模板标签文本匹配）
        # 需要从 ws 读取该行 A 列的实际标签文本
        actual_label = ws.cell(row=row_num, column=label_col).value
        actual_label_lower = (actual_label or "").lower().strip()

        item = label_to_item.get(actual_label_lower)
        if item:
            ws.cell(row=row_num, column=spec_col, value=item.get("specification", ""))
            ws.cell(row=row_num, column=result_col, value=item.get("result", ""))
            ws.cell(row=row_num, column=method_col, value=item.get("method", ""))
            mapped_keys.add(item_key)
            logger.info(f'[填充-数据] Row {row_num} ({actual_label}): '
                        f'spec={item.get("specification", "")}, result={item.get("result", "")}')
        else:
            # 尝试通过 item_key 匹配（兼容旧的 key-based 映射）
            for m in mapping:
                idx = m.get("test_item_index")
                if idx is None or idx >= len(test_items) or m.get("is_extra"):
                    continue
                # 如果 template_row_label 与 item_key 相关
                trl = (m.get("template_row_label") or "").lower().strip()
                if trl == actual_label_lower:
                    item = test_items[idx]
                    ws.cell(row=row_num, column=spec_col, value=item.get("specification", ""))
                    ws.cell(row=row_num, column=result_col, value=item.get("result", ""))
                    ws.cell(row=row_num, column=method_col, value=item.get("method", ""))
                    mapped_keys.add(item_key)
                    logger.info(f'[填充-数据] Row {row_num} ({actual_label}): matched via key')
                    break
            else:
                # PDF 中无此检测项 → 清空模板预填值
                ws.cell(row=row_num, column=spec_col, value="")
                ws.cell(row=row_num, column=result_col, value="")
                ws.cell(row=row_num, column=method_col, value="")
                logger.info(f'[填充-数据] Row {row_num} ({actual_label}): PDF 无对应数据，已清空')

    # --- 插入额外行（PDF 有值但模板无对应行） ---
    rows_inserted = 0
    if extra_items:
        # 按 insert_after_label 分组
        groups = defaultdict(list)
        for ei in extra_items:
            after_label = (ei.get("insert_after_label") or "").lower().strip()
            # 查找 after_label 对应的行号
            after_row = None
            if after_label:
                for item_key, row_num in layout.table_rows.items():
                    actual = ws.cell(row=row_num, column=label_col).value
                    if (actual or "").lower().strip() == after_label:
                        after_row = row_num
                        break
            groups[after_row].append(ei["item"])

        # 从行号最大的位置开始插入（避免偏移影响）
        for after_row in sorted(groups.keys(), key=lambda x: x if x is not None else -1, reverse=True):
            items = groups[after_row]
            if not items:
                continue
            if after_row is None:
                # 插在第一个数据行之前
                first_row = min(
                    (r for k, r in layout.table_rows.items() if k != 'assay'),
                    default=layout.data_header_row + 1 if layout.data_header_row else 10
                )
                insert_at = first_row
            else:
                insert_at = after_row + 1

            ws.insert_rows(insert_at, amount=len(items))
            for i, item in enumerate(items):
                row_num = insert_at + i
                ws.cell(row=row_num, column=label_col, value=item.get("name", ""))
                ws.cell(row=row_num, column=spec_col, value=item.get("specification", ""))
                ws.cell(row=row_num, column=result_col, value=item.get("result", ""))
                ws.cell(row=row_num, column=method_col, value=item.get("method", ""))
                logger.info(f'[填充-额外行] Row {row_num} ({item.get("name", "")})')

            rows_inserted += len(items)

    if rows_inserted:
        logger.info(f'[填充] 插入 {rows_inserted} 个额外行')

    # --- 填充 Packing & Storage ---
    # 因插入行导致位置下移
    if rows_inserted > 0:
        for attr in ('packing_position', 'storage_position'):
            pos = getattr(layout, attr, None)
            if pos and pos.row:
                pos.row += rows_inserted
                pos.cell_ref = None

    if layout.packing_position:
        ref = layout.packing_position.cell_ref
        if ref:
            ws[ref] = packing
        else:
            ws.cell(row=layout.packing_position.row,
                    column=layout.packing_position.col, value=packing)
        logger.info(f'[填充-包装] = {packing[:50]}' if packing else '[填充-包装] (空)')

    if layout.storage_position:
        ref = layout.storage_position.cell_ref
        if ref:
            ws[ref] = storage
        else:
            ws.cell(row=layout.storage_position.row,
                    column=layout.storage_position.col, value=storage)
        logger.info(f'[填充-存储] = {storage[:50]}' if storage else '[填充-存储] (空)')


def _fill_allergen_xlsx(ws, data: dict, layout: TemplateLayout):
    """填充 Allergen 模板 — 主要替换产品名"""
    product_name = data.get("header", {}).get("product_name", "")
    if product_name:
        pm = layout.header_fields.get('product_name')
        if pm and pm.value_pos and pm.value_pos.cell_ref:
            current = ws[pm.value_pos.cell_ref].value or ''
            if current.endswith(': ') or current.endswith(':'):
                ws[pm.value_pos.cell_ref] = current + product_name
            else:
                ws[pm.value_pos.cell_ref] = f'Product Name: {product_name}'
            logger.info(f'[填充-Allergen] 产品名: {product_name}')


# ============ Flow Chart 填充（XML 直接操作，保留 SmartArt） ============

def _fill_flowchart_xml(data: dict, layout: TemplateLayout, output_path: str):
    """使用 ZIP/XML 操作填充 Flow Chart 标题，保留 SmartArt 图表数据"""
    product_name = data.get("header", {}).get("product_name", "")
    if not product_name:
        logger.info('[填充-FlowChart] 产品名为空，跳过')
        return

    title_text = f'{product_name} Flow Chart'
    SHARED_STRINGS_PATH = 'xl/sharedStrings.xml'
    WORKBOOK_PATH = 'xl/workbook.xml'
    WORKBOOK_RELS_PATH = 'xl/_rels/workbook.xml.rels'
    A2_INDEX_PATTERN = re.compile(r'<c\s+r="A2"[^>]*t="s"[^>]*>\s*<v>(\d+)</v>', re.DOTALL)
    SI_PATTERN = re.compile(r'<si>.*?</si>', re.DOTALL)

    SHEET_KEYWORDS = [
        ('extract', '所有Extract'),
        ('juice concentrate', '浓缩汁'),
        ('concentrate', '浓缩汁'),
        ('freeze dried', '冻干粉'),
        ('freeze-dried', '冻干粉'),
        ('juice powder', '汁粉'),
    ]

    def _select_sheet_tab(product: str, sheet_names: list) -> int:
        pl = product.lower()
        for keyword, sheet_hint in SHEET_KEYWORDS:
            if keyword in pl:
                for idx, sn in enumerate(sheet_names):
                    if sheet_hint in sn:
                        return idx
        return 0

    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.xlsx')
    os.close(tmp_fd)
    try:
        with zipfile.ZipFile(output_path, 'r') as zin:
            wb_xml = zin.read(WORKBOOK_PATH).decode('UTF-8')
            sheet_entries = re.findall(
                r'<sheet\s+name="([^"]+)"\s+sheetId="\d+"\s+r:id="(rId\d+)"', wb_xml)
            sheet_names = [name for name, _ in sheet_entries]
            sheet_rids = {name: rid for name, rid in sheet_entries}

            rels_xml = zin.read(WORKBOOK_RELS_PATH).decode('UTF-8')
            rid_to_file = {}
            for rid, target in re.findall(r'<Relationship\s+Id="(rId\d+)"[^>]*Target="([^"]+)"', rels_xml):
                rid_to_file[rid] = target

            target_tab = _select_sheet_tab(product_name, sheet_names)
            target_sheet_name = sheet_names[target_tab]
            target_rid = sheet_rids[target_sheet_name]
            target_file = 'xl/' + rid_to_file[target_rid]

            sheet_xml = zin.read(target_file).decode('UTF-8')
            m = A2_INDEX_PATTERN.search(sheet_xml)
            if not m:
                logger.warning(f'[填充-FlowChart] 未找到 A2 共享字符串引用')
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
                return
            target_idx = int(m.group(1))

        NEW_SI = (
            '<si><r><rPr><b/><sz val="18"/><color rgb="FFFF0000"/>'
            '<rFont val="Times New Roman"/><family val="1"/></rPr>'
            f'<t>{product_name}</t></r><r><rPr><b/><sz val="18"/>'
            '<color theme="1"/><rFont val="Times New Roman"/><family val="1"/></rPr>'
            '<t xml:space="preserve"> Flow Chart</t></r></si>'
        )

        with zipfile.ZipFile(output_path, 'r') as zin, \
             zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                file_data = zin.read(item.filename)
                if item.filename == SHARED_STRINGS_PATH:
                    xml_str = file_data.decode('UTF-8')
                    si_elements = list(SI_PATTERN.finditer(xml_str))
                    if target_idx < len(si_elements):
                        old_si = si_elements[target_idx]
                        xml_str = xml_str[:old_si.start()] + NEW_SI + xml_str[old_si.end():]
                    file_data = xml_str.encode('UTF-8')
                elif item.filename == WORKBOOK_PATH:
                    xml_str = file_data.decode('UTF-8')
                    if 'activeTab=' in xml_str:
                        xml_str = re.sub(r'activeTab="\d+"', f'activeTab="{target_tab}"', xml_str)
                    else:
                        xml_str = xml_str.replace('<workbookView ', f'<workbookView activeTab="{target_tab}" ')
                    file_data = xml_str.encode('UTF-8')
                zout.writestr(item, file_data)
        shutil.move(tmp_path, output_path)
        logger.info(f'[填充-FlowChart] 标题: {title_text}')
    except Exception as e:
        logger.error(f'[填充-FlowChart] 失败: {e}')
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise


# ============ DOCX 填充 ============

def _fill_docx(data: dict, layout: TemplateLayout, template_path: str, output_path: str):
    """填充 DOCX 模板"""
    from docx import Document
    from docx.shared import RGBColor

    shutil.copy2(template_path, output_path)
    if sys.platform != 'win32':
        os.chmod(output_path, 0o644)

    doc = Document(output_path)
    header = data.get("header", {})
    product_name = header.get("product_name", "")
    country = header.get("country", "")
    botanical_name = header.get("botanical_name", "")

    logger.info(f'[DOCX填充] 模板类型: {layout.template_type}, 产品名: {product_name}')

    if layout.template_type == 'cs':
        _fill_cs_docx(doc, product_name, country)
    elif layout.template_type == 'nutrition':
        _fill_nutrition_docx(doc, product_name)
    elif layout.template_type == 'sds':
        _fill_sds_docx(doc, product_name, botanical_name)
    elif layout.template_type in ('composition_powder', 'composition_standard'):
        _fill_composition_docx(doc, product_name, data=data, layout=layout)
    else:
        _replace_product_name_in_paragraphs(doc, product_name)

    doc.save(output_path)
    logger.info(f'[DOCX填充] 已保存: {output_path}')


def _replace_in_paragraph(paragraph, old_text: str, new_text: str) -> bool:
    """在段落中替换文本（先单 run，再跨 run）"""
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    full_text = paragraph.text
    if old_text in full_text:
        pos = full_text.find(old_text)
        current_pos = 0
        for idx, run in enumerate(paragraph.runs):
            run_end = current_pos + len(run.text)
            if current_pos <= pos < run_end:
                offset = pos - current_pos
                remaining = run.text[offset:]
                if old_text.startswith(remaining):
                    run.text = run.text[:offset] + new_text
                    chars_consumed = len(remaining)
                    for next_run in paragraph.runs[idx + 1:]:
                        if chars_consumed >= len(old_text):
                            break
                        chars_to_remove = min(len(next_run.text), len(old_text) - chars_consumed)
                        next_run.text = next_run.text[chars_to_remove:]
                        chars_consumed += chars_to_remove
                    return True
            current_pos = run_end
    return False


def _replace_product_name_in_paragraphs(doc, product_name: str):
    """在段落中查找占位产品名并替换"""
    if not product_name:
        return
    for para in doc.paragraphs:
        m = re.search(r'our product[,\s]+(.+?)(?:\s+is[:\s]|\s*,)', para.text)
        if m:
            old_name = m.group(1).strip()
            if old_name and old_name != product_name:
                _replace_in_paragraph(para, old_name, product_name)


def _fill_cs_docx(doc, product_name: str, country: str):
    """Combined Statement"""
    _replace_product_name_in_paragraphs(doc, product_name)
    if country:
        for para in doc.paragraphs:
            m = re.match(r'(Country of Origin\s*[-–—]\s*)(.*)', para.text)
            if m:
                old_country = m.group(2).strip()
                if old_country and old_country != country:
                    _replace_in_paragraph(para, old_country, country)
                break


def _fill_nutrition_docx(doc, product_name: str):
    """Nutrition Info — 替换产品名并清空营养数据占位符"""
    _replace_product_name_in_paragraphs(doc, product_name)
    # 清空营养数据表中的占位符值（Row 1~11 的 C1 列），强制 Step 3.5 网络搜索
    for table in doc.tables:
        if len(table.rows) >= 2 and any(
            'item' in table.rows[0].cells[0].text.strip().lower()
            for _ in [1]
        ):
            for row in table.rows[1:]:
                if len(row.cells) >= 2:
                    val_cell = row.cells[1]
                    for para in val_cell.paragraphs:
                        for run in para.runs:
                            run.text = ''
            logger.info('[填充-Nutrition] 已清空营养数据占位符，需执行 Step 3.5 网络搜索')


def _fill_sds_docx(doc, product_name: str, botanical_name: str):
    """Safety Data Sheet"""
    # 替换表格中的产品名
    if product_name:
        for table in doc.tables:
            for row in table.rows:
                for ci, cell in enumerate(row.cells):
                    if 'product name' in cell.text.strip().lower():
                        if ci + 1 < len(row.cells):
                            val_cell = row.cells[ci + 1]
                            if val_cell.paragraphs and val_cell.paragraphs[0].runs:
                                val_cell.paragraphs[0].runs[0].text = product_name
                        break
    # 替换 Material Name
    if botanical_name:
        for table in doc.tables:
            for row in table.rows:
                for ci, cell in enumerate(row.cells):
                    if 'material name' in cell.text.strip().lower():
                        if ci + 1 < len(row.cells):
                            val_cell = row.cells[ci + 1]
                            if val_cell.paragraphs and val_cell.paragraphs[0].runs:
                                val_cell.paragraphs[0].runs[0].text = botanical_name
                        break
    # 去除黄色高亮
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.highlight_color is not None:
                            run.font.highlight_color = None


def _select_composition_page_from_json(data: dict, layout: TemplateLayout):
    """根据 JSON 数据选择 Composition 模板的目标页面索引"""
    assay = data.get("assay")
    if not assay:
        return None
    if layout.template_type == 'composition_powder':
        ratio_map = {'4:1': 0, '10:1': 1, '20:1': 2, '50:1': 3, '100:1': 4}
        for field in ('result', 'specification'):
            val = assay.get(field, '').strip()
            if val in ratio_map:
                return ratio_map[val]
        return None
    elif layout.template_type == 'composition_standard':
        result_text = assay.get('result', '')
        m = re.search(r'([\d.]+)\s*%', result_text)
        if m:
            pct = float(m.group(1))
            if pct < 50:
                return 0
            elif pct <= 80:
                return 1
            else:
                return 2
        return None
    return None


def _fill_composition_docx(doc, product_name: str, data: dict = None, layout: TemplateLayout = None):
    """Composition Statement"""
    from docx.shared import RGBColor

    if not product_name:
        return

    target_page = None
    if data and layout:
        target_page = _select_composition_page_from_json(data, layout)
        if target_page is not None:
            logger.info(f'[填充-Composition] 选择目标页面: {target_page}')

    # 替换段落中的产品名
    product_para_indices = []
    for i, para in enumerate(doc.paragraphs):
        m = re.search(r'our product,(.+?),\s*is', para.text)
        if m:
            product_para_indices.append(i)

    for idx, para_idx in enumerate(product_para_indices):
        if target_page is not None and idx != target_page:
            continue
        para = doc.paragraphs[para_idx]
        m = re.search(r'our product,(.+?),\s*is', para.text)
        if m:
            old_segment = m.group(1)
            if old_segment.strip() != product_name:
                old_full = m.group(0)
                new_full = f'our product, {product_name} ,is'
                if not _replace_in_paragraph(para, old_full, new_full):
                    new_text = para.text.replace(old_full, new_full)
                    if para.runs:
                        para.runs[0].text = new_text
                        for run in para.runs[1:]:
                            run.text = ''

    # 填充表格
    for ti, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue
        if target_page is not None and ti != target_page:
            continue
        row1 = table.rows[1]
        if len(row1.cells) >= 1:
            cell0 = row1.cells[0]
            # 遍历所有段落找到有内容的段落进行替换（处理前导空段落的情况）
            replaced = False
            for para in cell0.paragraphs:
                if para.runs and any(r.text.strip() for r in para.runs):
                    para.runs[0].text = product_name
                    for run in para.runs[1:]:
                        run.text = ''
                    replaced = True
                    break
            if not replaced and cell0.paragraphs and cell0.paragraphs[0].runs:
                cell0.paragraphs[0].runs[0].text = product_name
        if len(row1.cells) >= 2:
            cell1 = row1.cells[1]
            for para in cell1.paragraphs:
                text = para.text.strip()
                if not text or 'maltodextrin' in text.lower():
                    continue
                m = re.match(r'(\d+%)\s*(.*)', text)
                if m:
                    percentage = m.group(1)
                    new_text = f'{percentage} {product_name}'
                    if para.runs:
                        para.runs[0].text = new_text
                        for run in para.runs[1:]:
                            run.text = ''
                    break
    # 清除红色标记
    red_color = RGBColor(0xEE, 0x00, 0x00)
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.color and run.font.color.rgb == red_color:
                run.text = ''


# ============ CLI 入口 ============

def main():
    if len(sys.argv) < 4:
        print("用法: python3 fill_template.py <json_data_path> <template_path> <output_path>",
              file=sys.stderr)
        sys.exit(1)

    json_path = sys.argv[1]
    template_path = sys.argv[2]
    output_path = sys.argv[3]

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    version = data.get("schema_version", "")
    if version != "1.0":
        logger.warning(f'[填充] JSON schema 版本: {version}（期望 1.0）')

    fill_from_json(data, template_path, output_path)
    print(f'填充完成: {output_path}')


if __name__ == "__main__":
    main()
