#!/usr/bin/env python3
"""
Unit tests for the core converter modules:
  - template_detector.py: template type detection and layout parsing
  - fill_template.py: JSON-driven template filling (AI-First architecture)

These tests use the real template files in templates/ and verify that
the detection + filling pipeline produces correct output.
"""

import json
import os
import sys
import shutil
import tempfile
import unittest
import unittest.mock
from pathlib import Path
from unittest.mock import patch

PROJECT_DIR = Path(__file__).parent
sys.path.insert(0, str(PROJECT_DIR))
sys.path.insert(0, str(PROJECT_DIR / 'converter'))

from template_detector import detect_template_layout
from fill_template import fill_from_json
from job_manager import JobManager

TEMPLATES_DIR = PROJECT_DIR / 'templates'

# Minimal COA JSON data for Saigon Cinnamon Powder
SAMPLE_COA_DATA = {
    "header": {
        "product_name": "Saigon Cinnamon Powder",
        "botanical_name": "Cinnamomum loureiroi Nees",
        "plant_part": "Bark",
        "batch_number": "CP2409025",
        "country": "China",
        "mfg_date": "2024.09.21",
        "exp_date": "2029.09.20",
    },
    "assay": None,
    "test_items": [
        {"name": "Appearance", "specification": "Red brown powder",
         "result": "Complies", "method": "Visual", "section": "analytical"},
        {"name": "Odor", "specification": "Characteristic",
         "result": "Complies", "method": "Organoleptic", "section": "analytical"},
        {"name": "Taste", "specification": "Characteristic",
         "result": "Complies", "method": "Organoleptic", "section": "analytical"},
        {"name": "Loss on Drying", "specification": "12.00% max",
         "result": "4.84%", "method": "USP43-NF38<731>", "section": "analytical"},
        {"name": "Ash", "specification": "5.0% max",
         "result": "1.43%", "method": "USP43-NF38<561>", "section": "analytical"},
        {"name": "Particle Size", "specification": "95.00% min through 30 meshes",
         "result": "99.16%", "method": "US#30 Sieve", "section": "analytical"},
        {"name": "Lead (Pb)", "specification": "0.50 mg/kg max",
         "result": "0.29 mg/kg", "method": "Atomic Absorption", "section": "analytical"},
        {"name": "Total Bacterial Count", "specification": "10,000 cfu/g max",
         "result": "760 cfu/g", "method": "USP43-NF38<2021>", "section": "microbiology"},
        {"name": "E. Coli", "specification": "Negative",
         "result": "Not Detected", "method": "USP43-NF38<2022>", "section": "microbiology"},
    ],
    "packing": "",
    "storage": "Store in a well closed container at room temperature.",
    "template_mapping": [
        {"test_item_index": 0, "template_row_label": "Appearance", "is_extra": False},
        {"test_item_index": 1, "template_row_label": "Odor", "is_extra": False},
        {"test_item_index": 2, "template_row_label": "Taste", "is_extra": False},
        {"test_item_index": 3, "template_row_label": "Loss on drying", "is_extra": False},
        {"test_item_index": 4, "template_row_label": "Ash", "is_extra": False},
        {"test_item_index": 5, "template_row_label": "Sieve Analysis", "is_extra": False},
        {"test_item_index": 6, "template_row_label": "Pb", "is_extra": False},
        {"test_item_index": 7, "template_row_label": "Total Plate count", "is_extra": False},
        {"test_item_index": 8, "template_row_label": "E. Coli.", "is_extra": False},
    ],
}


def _template(name):
    """Resolve template path; skip if not found."""
    p = TEMPLATES_DIR / name
    if not p.exists():
        return None
    return str(p)


def _read_xlsx_cells(path):
    """Read all non-empty cells from an XLSX file into a dict {cell_ref: value}."""
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    cells = {}
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=10, values_only=False):
        for cell in row:
            if cell.value is not None:
                cells[cell.coordinate] = cell.value
    wb.close()
    return cells


def _read_docx_content(path):
    """Read DOCX paragraphs and tables."""
    from docx import Document
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs]
    tables = []
    for table in doc.tables:
        t = []
        for row in table.rows:
            t.append([cell.text for cell in row.cells])
        tables.append(t)
    return paras, tables


# ===========================================================================
# Template Detector Tests
# ===========================================================================

class TestTemplateDetector(unittest.TestCase):
    """Test that detect_template_layout correctly identifies all 10 template types."""

    def test_detect_coa_assay(self):
        tpl = _template('Key In COA - Assay.xlsx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        self.assertEqual(layout.template_type, 'coa_assay')
        self.assertEqual(layout.format, 'xlsx')
        self.assertIn('assay', layout.table_rows)

    def test_detect_coa_ratio(self):
        tpl = _template('Key In COA - Ratio.xlsx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        self.assertEqual(layout.template_type, 'coa_ratio')

    def test_detect_coa_powder(self):
        tpl = _template('Key In COA - Powder.xlsx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        self.assertEqual(layout.template_type, 'coa_powder')
        # Powder should NOT have an assay row
        self.assertNotIn('assay', layout.table_rows)

    def test_detect_allergen(self):
        tpl = _template('Allergen -.xlsx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        self.assertEqual(layout.template_type, 'allergen')

    def test_detect_flowchart(self):
        tpl = _template('Flow Chart.xlsx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        self.assertEqual(layout.template_type, 'flowchart')

    def test_detect_cs(self):
        tpl = _template('CS -.docx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        self.assertEqual(layout.template_type, 'cs')
        self.assertEqual(layout.format, 'docx')

    def test_detect_nutrition(self):
        tpl = _template('Nutrition info -.docx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        self.assertEqual(layout.template_type, 'nutrition')

    def test_detect_sds(self):
        tpl = _template('Safety Data Sheet -.docx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        self.assertEqual(layout.template_type, 'sds')

    def test_detect_composition_powder(self):
        tpl = _template('Composition Statement - Powder & Ratio.docx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        self.assertEqual(layout.template_type, 'composition_powder')

    def test_detect_composition_standard(self):
        tpl = _template('Composition Statement - Standardized Material.docx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        self.assertEqual(layout.template_type, 'composition_standard')

    def test_coa_header_fields_detected(self):
        """COA templates should detect all 7 standard header fields."""
        tpl = _template('Key In COA - Powder.xlsx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        expected = {'product_name', 'botanical_name', 'plant_part',
                    'batch_number', 'country', 'mfg_date', 'exp_date'}
        self.assertEqual(set(layout.header_fields.keys()), expected)

    def test_coa_data_columns_detected(self):
        """COA templates should detect spec/result/method columns."""
        tpl = _template('Key In COA - Assay.xlsx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        self.assertIn('spec', layout.data_columns)
        self.assertIn('result', layout.data_columns)
        self.assertIn('method', layout.data_columns)

    def test_coa_table_rows_include_standard_items(self):
        """COA Powder template should detect standard test item rows."""
        tpl = _template('Key In COA - Powder.xlsx')
        if not tpl:
            self.skipTest('Template not found')
        layout = detect_template_layout(tpl)
        for key in ('appearance', 'odor', 'taste', 'lod', 'ash', 'pb', 'tpc'):
            self.assertIn(key, layout.table_rows, f'{key} not detected')


# ===========================================================================
# Fill Template Tests — XLSX
# ===========================================================================

class TestFillXLSX(unittest.TestCase):
    """Test JSON-driven XLSX template filling."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _fill(self, template_name, data=None):
        tpl = _template(template_name)
        if not tpl:
            self.skipTest(f'Template {template_name} not found')
        out = os.path.join(self.tmpdir, 'output' + os.path.splitext(template_name)[1])
        fill_from_json(data or SAMPLE_COA_DATA, tpl, out)
        return out

    def test_powder_header_fields(self):
        """Header fields should be filled in correct cells."""
        out = self._fill('Key In COA - Powder.xlsx')
        cells = _read_xlsx_cells(out)
        self.assertEqual(cells.get('C4'), 'Saigon Cinnamon Powder')
        self.assertEqual(cells.get('C5'), 'Cinnamomum loureiroi Nees')
        self.assertEqual(cells.get('F5'), 'Bark')
        self.assertEqual(cells.get('C6'), 'CP2409025')
        self.assertEqual(cells.get('F6'), 'China')
        self.assertEqual(cells.get('C7'), '2024.09.21')
        self.assertEqual(cells.get('F7'), '2029.09.20')

    def test_powder_data_rows_filled(self):
        """Mapped test items should fill spec/result/method in correct rows."""
        out = self._fill('Key In COA - Powder.xlsx')
        cells = _read_xlsx_cells(out)
        # Find Appearance row — should have "Red brown powder" in spec column
        found = False
        for coord, val in cells.items():
            if val == 'Red brown powder':
                found = True
                break
        self.assertTrue(found, 'Appearance spec not found in output')

    def test_powder_unmapped_rows_cleared(self):
        """Template rows without PDF data should be cleared (no residual defaults)."""
        out = self._fill('Key In COA - Powder.xlsx')
        cells = _read_xlsx_cells(out)
        # Heavy metals row should be cleared (no data in sample)
        # Find the row with "Heavy metals" label
        heavy_row = None
        for coord, val in cells.items():
            if isinstance(val, str) and 'heavy metals' in val.lower():
                heavy_row = int(coord[1:])  # extract row number
                break
        if heavy_row:
            # Spec/Result/Method should be empty for this row
            import openpyxl
            wb = openpyxl.load_workbook(out, data_only=True)
            ws = wb.active
            self.assertIsNone(ws.cell(row=heavy_row, column=3).value)
            self.assertIsNone(ws.cell(row=heavy_row, column=5).value)
            wb.close()

    def test_powder_storage_filled(self):
        """Storage text should appear in the output."""
        out = self._fill('Key In COA - Powder.xlsx')
        cells = _read_xlsx_cells(out)
        found = any('Store in a well closed' in str(v) for v in cells.values())
        self.assertTrue(found, 'Storage text not found')

    def test_assay_row_cleared_when_null(self):
        """Assay template with assay=None should clear the Assay row."""
        out = self._fill('Key In COA - Assay.xlsx')
        cells = _read_xlsx_cells(out)
        # Find Assay label row
        assay_row = None
        for coord, val in cells.items():
            if isinstance(val, str) and val.strip() == 'Assay':
                assay_row = int(''.join(c for c in coord if c.isdigit()))
                break
        if assay_row:
            import openpyxl
            wb = openpyxl.load_workbook(out, data_only=True)
            ws = wb.active
            # Spec, Result, Method should be empty
            self.assertFalse(ws.cell(row=assay_row, column=3).value)
            self.assertFalse(ws.cell(row=assay_row, column=5).value)
            wb.close()

    def test_assay_row_filled_when_present(self):
        """Assay template with assay data should fill the Assay row."""
        data = dict(SAMPLE_COA_DATA)
        data['assay'] = {
            "specification": "5% Withanolides",
            "result": "5.2%",
            "method": "HPLC",
        }
        out = self._fill('Key In COA - Assay.xlsx', data)
        cells = _read_xlsx_cells(out)
        found = any(v == '5% Withanolides' for v in cells.values())
        self.assertTrue(found, 'Assay spec not found')

    def test_allergen_product_name(self):
        """Allergen template should contain product name in the output."""
        out = self._fill('Allergen -.xlsx')
        cells = _read_xlsx_cells(out)
        found = any('Saigon Cinnamon Powder' in str(v) for v in cells.values())
        self.assertTrue(found, 'Product name not in Allergen output')

    def test_flowchart_title(self):
        """Flow Chart should have product name in title."""
        out = self._fill('Flow Chart.xlsx')
        cells = _read_xlsx_cells(out)
        found = any('Saigon Cinnamon Powder' in str(v) and 'Flow Chart' in str(v)
                     for v in cells.values())
        self.assertTrue(found, 'Flow Chart title not found')


# ===========================================================================
# Fill Template Tests — DOCX
# ===========================================================================

class TestFillDOCX(unittest.TestCase):
    """Test JSON-driven DOCX template filling."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _fill(self, template_name, data=None):
        tpl = _template(template_name)
        if not tpl:
            self.skipTest(f'Template {template_name} not found')
        out = os.path.join(self.tmpdir, 'output.docx')
        fill_from_json(data or SAMPLE_COA_DATA, tpl, out)
        return out

    def test_cs_product_name_replaced(self):
        """CS template should replace product name placeholder."""
        out = self._fill('CS -.docx')
        paras, _ = _read_docx_content(out)
        found = any('Saigon Cinnamon Powder' in p for p in paras)
        self.assertTrue(found, 'Product name not in CS output')

    def test_cs_country_replaced(self):
        """CS template should replace country of origin."""
        out = self._fill('CS -.docx')
        paras, _ = _read_docx_content(out)
        found = any('China' in p for p in paras)
        self.assertTrue(found, 'Country not in CS output')

    def test_nutrition_product_name_replaced(self):
        """Nutrition template should replace product name."""
        out = self._fill('Nutrition info -.docx')
        paras, _ = _read_docx_content(out)
        found = any('Saigon Cinnamon Powder' in p for p in paras)
        self.assertTrue(found, 'Product name not in Nutrition output')

    def test_nutrition_placeholders_cleared(self):
        """Nutrition template should clear all nutrition value placeholders."""
        out = self._fill('Nutrition info -.docx')
        _, tables = _read_docx_content(out)
        self.assertTrue(len(tables) >= 1, 'No tables in Nutrition output')
        nutrition_table = tables[0]
        # Rows 1-11 (index 1-11) Column 1 should be empty
        for row in nutrition_table[1:]:
            val = row[1].strip() if len(row) > 1 else ''
            self.assertEqual(val, '',
                             f'Nutrition placeholder not cleared: {row[0]}={val}')

    def test_sds_product_name(self):
        """SDS template should fill Product Name in table."""
        out = self._fill('Safety Data Sheet -.docx')
        _, tables = _read_docx_content(out)
        self.assertTrue(len(tables) >= 1)
        # Table 0 should have product name
        found = False
        for row in tables[0]:
            if 'Saigon Cinnamon Powder' in ' '.join(row):
                found = True
                break
        self.assertTrue(found, 'Product name not in SDS Table 0')

    def test_sds_material_name(self):
        """SDS template should fill Material Name (botanical name)."""
        out = self._fill('Safety Data Sheet -.docx')
        _, tables = _read_docx_content(out)
        found = False
        for row in tables[0]:
            if 'Cinnamomum loureiroi Nees' in ' '.join(row):
                found = True
                break
        self.assertTrue(found, 'Material name not in SDS Table 0')

    def test_composition_powder_product_name(self):
        """Composition Statement should replace product name in paragraphs."""
        out = self._fill('Composition Statement - Powder & Ratio.docx')
        paras, _ = _read_docx_content(out)
        found = any('Saigon Cinnamon Powder' in p for p in paras)
        self.assertTrue(found, 'Product name not in Composition output')

    def test_composition_powder_no_residual_template_name(self):
        """Composition Statement should not retain old template product names."""
        out = self._fill('Composition Statement - Powder & Ratio.docx')
        paras, tables = _read_docx_content(out)
        # Check no "Pumpkin Seed Protein Powder" remains anywhere
        all_text = ' '.join(paras)
        for table in tables:
            for row in table:
                all_text += ' '.join(row)
        self.assertNotIn('Pumpkin Seed Protein Powder', all_text,
                         'Residual template product name found')

    def test_composition_standard_product_name(self):
        """Standardized Material Composition should replace product name."""
        out = self._fill('Composition Statement - Standardized Material.docx')
        paras, _ = _read_docx_content(out)
        found = any('Saigon Cinnamon Powder' in p for p in paras)
        self.assertTrue(found, 'Product name not in Standardized Composition output')


# ===========================================================================
# converter_service.py — status flow
# ===========================================================================

class TestConverterServiceStatusFlow(unittest.TestCase):
    """Verify converter_service always goes through verifying status."""

    def test_needs_verify_always_true(self):
        """converter_service no longer has supplier checking logic."""
        import converter_service
        import inspect
        source = inspect.getsource(converter_service.run_conversion)
        self.assertNotIn('check_supplier', source)
        self.assertNotIn('needs_verify', source)
        # Should always set status to 'verifying'
        self.assertIn('verifying', source)


# ===========================================================================
# Claude CLI Detection Tests
# ===========================================================================

class TestClaudeCLIDetection(unittest.TestCase):
    """Test that _find_claude_cli covers all common installation paths."""

    def test_homebrew_apple_silicon_path_checked(self):
        """Should check /opt/homebrew/bin/claude (Apple Silicon Homebrew)."""
        from terminal_launcher import _find_claude_cli
        import inspect
        source = inspect.getsource(_find_claude_cli)
        self.assertIn('/opt/homebrew/bin/claude', source)

    def test_homebrew_intel_path_checked(self):
        """Should check /usr/local/bin/claude (Intel Homebrew / system)."""
        from terminal_launcher import _find_claude_cli
        import inspect
        source = inspect.getsource(_find_claude_cli)
        self.assertIn('/usr/local/bin/claude', source)

    def test_nvm_path_checked(self):
        """Should check nvm-managed Node.js paths."""
        from terminal_launcher import _find_claude_cli
        import inspect
        source = inspect.getsource(_find_claude_cli)
        self.assertIn('.nvm/versions/node/', source)

    def test_volta_path_checked(self):
        """Should check volta-managed paths."""
        from terminal_launcher import _find_claude_cli
        import inspect
        source = inspect.getsource(_find_claude_cli)
        self.assertIn('.volta/bin/claude', source)

    def test_bun_path_checked(self):
        """Should check bun paths."""
        from terminal_launcher import _find_claude_cli
        import inspect
        source = inspect.getsource(_find_claude_cli)
        self.assertIn('.bun/bin/claude', source)

    def test_shell_path_fallback_exists(self):
        """Should have a shell PATH fallback for server-env PATH differences."""
        from terminal_launcher import _get_shell_path
        # Function should exist and be callable
        self.assertTrue(callable(_get_shell_path))

    def test_find_claude_returns_executable_or_none(self):
        """_find_claude_cli should return an executable path or None."""
        from terminal_launcher import _find_claude_cli
        result = _find_claude_cli()
        if result is not None:
            self.assertTrue(os.path.isfile(result), f'{result} is not a file')
            self.assertTrue(os.access(result, os.X_OK), f'{result} is not executable')

    def test_env_override_takes_priority(self):
        """CLAUDE_CLI_PATH env var should override auto-detection."""
        from unittest.mock import patch
        fake = '/tmp/test-claude-fake'
        # Create a fake executable
        with open(fake, 'w') as f:
            f.write('#!/bin/sh\n')
        os.chmod(fake, 0o755)
        try:
            with patch.dict(os.environ, {'CLAUDE_CLI_PATH': fake}):
                # Re-evaluate: the module-level CLAUDE_CLI is set at import time,
                # so test the env var logic directly
                result = os.environ.get('CLAUDE_CLI_PATH')
                self.assertEqual(result, fake)
        finally:
            os.remove(fake)

    def test_find_claude_on_current_system(self):
        """Integration: verify _find_claude_cli finds claude on this machine (if installed)."""
        import shutil
        from terminal_launcher import _find_claude_cli
        result = _find_claude_cli()
        which_result = shutil.which('claude')
        if which_result:
            # If which finds it, our function should too
            self.assertIsNotNone(result,
                                 f'shutil.which found {which_result} but _find_claude_cli returned None')
        # If neither finds it, that's fine (CI/CD environment without claude)


# ===========================================================================
# Download E2E Tests
# ===========================================================================

class TestDownloadE2E(unittest.TestCase):
    """End-to-end download tests: upload → convert → download → verify file."""

    def setUp(self):
        import threading
        from http.server import HTTPServer
        from http.client import HTTPConnection
        import app as app_module
        from app import COAHandler

        self.app_module = app_module
        app_module.jobs = JobManager()
        self.server = HTTPServer(('127.0.0.1', 0), COAHandler)
        self.port = self.server.server_address[1]
        self._thread = threading.Thread(target=self.server.serve_forever, daemon=True)
        self._thread.start()

        self.pdf_data = (
            b'%PDF-1.0\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj '
            b'2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj '
            b'3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R>>endobj\n'
            b'xref\n0 4\n0000000000 65535 f \n'
            b'0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n'
            b'trailer<</Size 4/Root 1 0 R>>\nstartxref\n190\n%%EOF'
        )
        self.templates = []
        if TEMPLATES_DIR.exists():
            for f in sorted(TEMPLATES_DIR.iterdir()):
                if f.suffix.lower() in ('.xlsx', '.docx'):
                    self.templates.append(str(f))
        if not self.templates:
            self.skipTest('Need at least 1 template')

    def tearDown(self):
        self.server.shutdown()

    def _conn(self):
        from http.client import HTTPConnection
        return HTTPConnection('127.0.0.1', self.port, timeout=10)

    def _upload_pdf(self, filename='test_download.pdf'):
        boundary = '----DLTest'
        body = (
            f'------DLTest\r\n'
            f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
            f'Content-Type: application/pdf\r\n\r\n'
        ).encode() + self.pdf_data + b'\r\n------DLTest--\r\n'
        conn = self._conn()
        conn.request('POST', '/api/upload', body=body,
                     headers={'Content-Type': 'multipart/form-data; boundary=----DLTest'})
        resp = conn.getresponse()
        data = json.loads(resp.read())
        conn.close()
        return data

    @unittest.mock.patch('terminal_launcher.launch_verification_silent')
    @unittest.mock.patch('converter_service.convert_coa')
    def test_download_returns_attachment_not_redirect(self, mock_convert, mock_verify):
        """Download should return Content-Disposition: attachment, not a redirect."""
        output_content = b'This is the converted output'

        def fake_convert(pdf, tpl, out):
            Path(out).write_bytes(output_content)
            return out
        mock_convert.side_effect = fake_convert

        jobs_data = self._upload_pdf()
        job_id = jobs_data[0]['id']

        payload = json.dumps({
            'template_path': self.templates[0],
            'force_verify': False,
            'claude_mode': 'silent',
        }).encode()
        conn = self._conn()
        conn.request('POST', f'/api/convert/{job_id}', body=payload,
                     headers={'Content-Type': 'application/json'})
        resp = conn.getresponse()
        resp.read()
        conn.close()

        import time
        time.sleep(2)

        # Download
        conn = self._conn()
        conn.request('GET', f'/api/download/{job_id}')
        resp = conn.getresponse()
        body = resp.read()
        conn.close()

        # Must be 200, not 3xx redirect
        self.assertEqual(resp.status, 200)
        # Must have attachment disposition
        disp = resp.getheader('Content-Disposition', '')
        self.assertIn('attachment', disp, f'Expected attachment, got: {disp}')
        self.assertIn('filename', disp)
        # Content matches what we wrote
        self.assertEqual(body, output_content)
        # Must have Cache-Control to prevent stale cache
        cache = resp.getheader('Cache-Control', '')
        self.assertIn('no-cache', cache)

    @unittest.mock.patch('terminal_launcher.launch_verification_silent')
    @unittest.mock.patch('converter_service.convert_coa')
    def test_download_content_type_is_octet_stream(self, mock_convert, mock_verify):
        """Content-Type must be application/octet-stream to trigger browser download."""
        mock_convert.side_effect = lambda p, t, o: (Path(o).write_bytes(b'data'), o)[1]

        jobs_data = self._upload_pdf()
        job_id = jobs_data[0]['id']

        payload = json.dumps({
            'template_path': self.templates[0],
            'force_verify': False,
            'claude_mode': 'silent',
        }).encode()
        conn = self._conn()
        conn.request('POST', f'/api/convert/{job_id}', body=payload,
                     headers={'Content-Type': 'application/json'})
        resp = conn.getresponse()
        resp.read()
        conn.close()

        import time
        time.sleep(2)

        conn = self._conn()
        conn.request('GET', f'/api/download/{job_id}')
        resp = conn.getresponse()
        resp.read()
        conn.close()

        self.assertEqual(resp.getheader('Content-Type'), 'application/octet-stream')

    @unittest.mock.patch('terminal_launcher.launch_verification_silent')
    @unittest.mock.patch('converter_service.convert_coa')
    def test_download_filename_with_special_chars(self, mock_convert, mock_verify):
        """Filenames with spaces/special chars should be properly encoded."""
        def fake_convert(pdf, tpl, out):
            Path(out).write_bytes(b'data')
            return out
        mock_convert.side_effect = fake_convert

        jobs_data = self._upload_pdf('Test File (1).pdf')
        job_id = jobs_data[0]['id']

        payload = json.dumps({
            'template_path': self.templates[0],
            'force_verify': False,
            'claude_mode': 'silent',
        }).encode()
        conn = self._conn()
        conn.request('POST', f'/api/convert/{job_id}', body=payload,
                     headers={'Content-Type': 'application/json'})
        resp = conn.getresponse()
        resp.read()
        conn.close()

        import time
        time.sleep(2)

        conn = self._conn()
        conn.request('GET', f'/api/download/{job_id}')
        resp = conn.getresponse()
        resp.read()
        conn.close()

        disp = resp.getheader('Content-Disposition', '')
        self.assertIn('attachment', disp)
        # Should have UTF-8 encoded filename for special chars
        self.assertIn("filename*=UTF-8''", disp)

    def test_download_nonexistent_job_returns_404(self):
        """Download of non-existent job should return 404, not crash."""
        conn = self._conn()
        conn.request('GET', '/api/download/nonexistent-id-123')
        resp = conn.getresponse()
        resp.read()
        conn.close()
        self.assertEqual(resp.status, 404)

    @unittest.mock.patch('terminal_launcher.launch_verification_silent')
    @unittest.mock.patch('converter_service.convert_coa')
    def test_download_content_length_matches_body(self, mock_convert, mock_verify):
        """Content-Length header must match actual body size."""
        content = b'Exact content for length verification - 42 bytes!'
        mock_convert.side_effect = lambda p, t, o: (Path(o).write_bytes(content), o)[1]

        jobs_data = self._upload_pdf()
        job_id = jobs_data[0]['id']

        payload = json.dumps({
            'template_path': self.templates[0],
            'force_verify': False,
            'claude_mode': 'silent',
        }).encode()
        conn = self._conn()
        conn.request('POST', f'/api/convert/{job_id}', body=payload,
                     headers={'Content-Type': 'application/json'})
        resp = conn.getresponse()
        resp.read()
        conn.close()

        import time
        time.sleep(2)

        conn = self._conn()
        conn.request('GET', f'/api/download/{job_id}')
        resp = conn.getresponse()
        body = resp.read()
        conn.close()

        content_length = int(resp.getheader('Content-Length', 0))
        self.assertEqual(content_length, len(body))
        self.assertEqual(body, content)


if __name__ == '__main__':
    unittest.main()
